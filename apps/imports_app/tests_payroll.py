from io import BytesIO
from django.test import TestCase
from docx import Document
from apps.imports_app.payroll import import_payroll
from apps.roster.models import PayrollRecord

class PayrollImportTests(TestCase):
    def test_word_payroll_import(self):
        doc=Document()
        doc.add_paragraph("Hours for week ending: 02/08/2026")
        table=doc.add_table(rows=1, cols=4)
        table.rows[0].cells[0].text="Name"
        table.rows[0].cells[1].text="Ordinary Hours"
        table.rows[0].cells[2].text="Sunday Hours"
        table.rows[0].cells[3].text="Notes"
        row=table.add_row().cells
        row[0].text="Catherine Kirby"; row[1].text="28.50"; row[2].text="6.00"
        stream=BytesIO(); doc.save(stream); stream.seek(0); stream.name="wages.docx"
        week,count,issues=import_payroll(stream)
        self.assertEqual(str(week.week_end), "2026-08-02")
        self.assertEqual(count,1)
        self.assertEqual(float(PayrollRecord.objects.get().total_hours),34.5)
