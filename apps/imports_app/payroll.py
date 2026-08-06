import csv
from datetime import datetime
from decimal import Decimal, InvalidOperation
from io import TextIOWrapper
import re

from django.db import transaction
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from apps.employees.models import Department
from apps.roster.models import PayrollRecord, PayrollWeek
from .services import _employee_for_name, _normalise_name, _text

DATE_RE = re.compile(r"(?:week ending|hours for week ending|w[\s/]*e)[:\s]*(\d{1,2})[./-](\d{1,2})[./-](\d{2,4})", re.I)


def _hours(value):
    text = _text(value).strip()
    if not text or text.lower() in {"-", "--", "---", "----", "salary", "n/a"}:
        return Decimal("0")
    match = re.search(r"-?\d+(?:[.,]\d+)?", text)
    if not match:
        return Decimal("0")
    try:
        return Decimal(match.group(0).replace(",", "."))
    except InvalidOperation:
        return Decimal("0")


def _week_end(text):
    match = DATE_RE.search(text)
    if not match:
        return None
    day, month, year = map(int, match.groups())
    if year < 100:
        year += 2000
    return datetime(year, month, day).date()


def _docx_rows(file_obj):
    doc = Document(file_obj)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    week_end = _week_end(full_text)
    rows = []
    for table in doc.tables:
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
    if not rows:
        # Accountant-style Word reports are sometimes plain lines rather than tables.
        for line in full_text.splitlines():
            match = re.match(r"^\s*\d+\s+(.+?)\s+(Salary|[-–—]+|\d+(?:[.,]\d+)?)\s*(Salary|[-–—]+|\d+(?:[.,]\d+)?)?\s*$", line)
            if match:
                rows.append([match.group(1), match.group(2), match.group(3) or ""])
    return week_end, rows



def _plain_text_rows(full_text):
    """Convert accountant-style text reports into payroll-style rows."""
    rows = []
    pending_name = None

    for raw_line in full_text.splitlines():
        line = re.sub(r"\s+", " ", raw_line).strip()
        if not line:
            continue

        # Ignore document headings and column labels.
        lower = line.lower()
        if (
            "hours for week ending" in lower
            or lower.startswith("week ending")
            or (
                "name" in lower
                and "ordinary" in lower
                and "hours" in lower
            )
            or lower in {"sunday hours notes", "sunday hours", "notes"}
        ):
            continue

        # Typical line:
        # 9 Catherine Kirby 28.50 6.00
        match = re.match(
            r"^\s*(?:\d+\s+)?"
            r"(?P<name>[A-Za-zÀ-ž'’.\- ]+?)\s+"
            r"(?P<ordinary>Salary|[-–—]+|\d+(?:[.,]\d+)?)"
            r"(?:\s+(?P<sunday>Salary|[-–—]+|\d+(?:[.,]\d+)?))?"
            r"(?:\s+(?P<overtime>\d+(?:[.,]\d+)?))?"
            r"(?:\s+(?P<notes>.*))?$",
            line,
            re.I,
        )
        if match:
            name = match.group("name").strip()
            # Avoid treating footer prose as employee names.
            if len(name.split()) <= 5:
                rows.append(
                    [
                        name,
                        match.group("ordinary") or "",
                        match.group("sunday") or "",
                        match.group("overtime") or "",
                        match.group("notes") or "",
                    ]
                )
                pending_name = None
                continue

        # Some PDF converters split an employee row over two lines.
        if re.match(r"^(?:\d+\s+)?[A-Za-zÀ-ž'’.\- ]+$", line):
            candidate = re.sub(r"^\d+\s+", "", line).strip()
            if 1 <= len(candidate.split()) <= 5:
                pending_name = candidate
            continue

        if pending_name:
            hours_match = re.match(
                r"^(Salary|[-–—]+|\d+(?:[.,]\d+)?)"
                r"(?:\s+(Salary|[-–—]+|\d+(?:[.,]\d+)?))?"
                r"(?:\s+(\d+(?:[.,]\d+)?))?"
                r"(?:\s+(.*))?$",
                line,
                re.I,
            )
            if hours_match:
                rows.append(
                    [
                        pending_name,
                        hours_match.group(1) or "",
                        hours_match.group(2) or "",
                        hours_match.group(3) or "",
                        hours_match.group(4) or "",
                    ]
                )
                pending_name = None

    return rows


def _pdf_rows(file_obj):
    reader = PdfReader(file_obj)
    page_text = []

    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            page_text.append(text)

    full_text = "\n".join(page_text).strip()
    if not full_text:
        raise ValueError(
            "This PDF contains no selectable text. "
            "Use the original digital PDF, Word, Excel or CSV file."
        )

    week_end = _week_end(full_text)
    rows = _plain_text_rows(full_text)

    if not rows:
        raise ValueError(
            "The PDF text was readable, but no payroll employee rows "
            "could be recognised."
        )

    return week_end, rows


def _csv_rows(file_obj):
    wrapper = TextIOWrapper(file_obj, encoding="utf-8-sig", errors="replace")
    rows = list(csv.reader(wrapper))
    wrapper.detach()
    week_end = _week_end("\n".join(",".join(row) for row in rows[:15]))
    return week_end, rows


def _xlsx_rows(file_obj):
    workbook = load_workbook(file_obj, data_only=True)
    sheet = workbook.active
    rows = [[_text(v) for v in row] for row in sheet.iter_rows(values_only=True)]
    week_end = _week_end("\n".join(" ".join(row) for row in rows[:15]))
    return week_end, rows


def _locate_columns(rows):
    for index, row in enumerate(rows[:20]):
        values = [_text(v).lower().strip() for v in row]
        joined = " ".join(values)
        if "name" not in joined and "employee" not in joined:
            continue
        def find(*terms):
            return next((i for i, value in enumerate(values) if any(term in value for term in terms)), None)
        return index, {
            "name": find("name", "employee"),
            "ordinary": find("ordinary", "normal"),
            "sunday": find("sunday"),
            "overtime": find("overtime", "ot hours"),
            "total": find("total paid", "total hours", "paid hours"),
            "notes": find("notes", "comment"),
        }
    return None, None


def _fallback_docx_records(rows):
    records = []
    for row_number, row in enumerate(rows, 1):
        cells = [_text(v).strip() for v in row if _text(v).strip()]
        if not cells:
            continue
        if cells[0].isdigit():
            cells = cells[1:]
        if len(cells) < 2:
            continue
        name = _normalise_name(cells[0])
        if not name or name.lower() in {"name", "ordinary hours"}:
            continue
        ordinary = _hours(cells[1])
        sunday = _hours(cells[2]) if len(cells) > 2 else Decimal("0")
        notes = " ".join(cells[3:]) if len(cells) > 3 else ""
        records.append((row_number, name, ordinary, sunday, Decimal("0"), ordinary+sunday, notes))
    return records


def _normal_records(rows):
    header_index, columns = _locate_columns(rows)
    if header_index is None or columns["name"] is None:
        return _fallback_docx_records(rows)
    records=[]
    for row_number, row in enumerate(rows[header_index+1:], header_index+2):
        def cell(index):
            return row[index] if index is not None and index < len(row) else ""
        name=_normalise_name(cell(columns["name"]))
        if not name or name.lower() in {"total", "totals"}:
            continue
        ordinary=_hours(cell(columns["ordinary"]))
        sunday=_hours(cell(columns["sunday"]))
        overtime=_hours(cell(columns["overtime"]))
        total=_hours(cell(columns["total"]))
        if total == 0:
            total=ordinary+sunday+overtime
        notes=_text(cell(columns["notes"]))
        records.append((row_number,name,ordinary,sunday,overtime,total,notes))
    return records


@transaction.atomic
def import_payroll(file_obj):
    filename = getattr(file_obj, "name", "payroll").lower()
    if filename.endswith(".pdf"):
        week_end, rows = _pdf_rows(file_obj)
    elif filename.endswith(".docx"):
        week_end, rows = _docx_rows(file_obj)
    elif filename.endswith(".xlsx"):
        week_end, rows = _xlsx_rows(file_obj)
    elif filename.endswith(".csv"):
        week_end, rows = _csv_rows(file_obj)
    else:
        raise ValueError("Use a PDF, CSV, Excel or Word payroll file.")
    if week_end is None:
        raise ValueError("No readable payroll week-ending date was found.")

    payroll_week, _ = PayrollWeek.objects.update_or_create(
        week_end=week_end,
        defaults={"source_name": getattr(file_obj, "name", "")},
    )
    issues=[]; imported=0
    for row_number,name,ordinary,sunday,overtime,total,notes in _normal_records(rows):
        # Keep salary-only and empty rows visible as issues, but don't invent a target.
        if total <= 0:
            if notes or ordinary == 0:
                issues.append({"row":row_number,"name":name,"message":"No hourly total; ignored for target hours."})
            continue
        employee=_employee_for_name(name, Department.RESTAURANT)
        PayrollRecord.objects.update_or_create(
            payroll_week=payroll_week,
            employee=employee,
            defaults={
                "ordinary_hours":ordinary,
                "sunday_hours":sunday,
                "overtime_hours":overtime,
                "total_hours":total,
                "notes":notes,
                "source_row":row_number,
            },
        )
        imported += 1
    return payroll_week, imported, issues
